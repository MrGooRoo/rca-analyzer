"""
Извлечение текстового содержимого из DOCX-файлов.

Использует python-docx для парсинга.
Возвращает полный текст документа (абзацы + таблицы).
"""

from __future__ import annotations

import io
import logging

from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Извлечь весь текст из DOCX-файла.

    Включает текст абзацев и ячеек таблиц.
    Возвращает объединённый текст (разделённый переносами строк).
    """
    doc = Document(io.BytesIO(file_bytes))

    parts: list[str] = []

    # Абзацы
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Таблицы
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append(" | ".join(row_texts))

    full_text = "\n".join(parts)

    logger.info(
        "[DocxExtractor] Извлечено %d символов из DOCX (%d абзацев, %d таблиц)",
        len(full_text),
        len(doc.paragraphs),
        len(doc.tables),
    )

    return full_text
