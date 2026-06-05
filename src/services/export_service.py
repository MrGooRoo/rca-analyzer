"""
Export Service — генерация DOCX-отчёта по результату RCA-анализа.

Поддерживает все 5 методологий. Для bowtie определяет роль узла
по category (BOWTIE:THREAT, BOWTIE:PREVENTION, BOWTIE:MITIGATION,
BOWTIE:CONSEQUENCE, BOWTIE:HAZARD, BOWTIE:TOP_EVENT).

Зависимость: python-docx (уже в pyproject.toml как python-docx)
"""

from __future__ import annotations

import io
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.domain.models import RCAResult, MethodologyType


# ---------------------------------------------------------------------------
# Цветовая палитра (RGB)
# ---------------------------------------------------------------------------

_C_TITLE     = RGBColor(0x1a, 0x1a, 0x2e)   # почти чёрный
_C_PRIMARY   = RGBColor(0x01, 0x69, 0x6f)   # teal — основной акцент
_C_RED       = RGBColor(0xf7, 0x6f, 0x6f)   # корневые причины / угрозы
_C_ORANGE    = RGBColor(0xf7, 0xb9, 0x55)   # вносящие факторы / барьеры
_C_BLUE      = RGBColor(0x4f, 0x8e, 0xf7)   # непосред. причины / последствия
_C_GREEN     = RGBColor(0x3e, 0xcf, 0x8e)   # рекомендации
_C_MUTED     = RGBColor(0x7a, 0x79, 0x74)   # вспомогательный текст

_PRIORITY_COLORS = {
    "high":   _C_RED,
    "medium": _C_ORANGE,
    "low":    _C_GREEN,
}

METHODOLOGY_LABELS = {
    MethodologyType.FIVE_WHY:     "5 Почему",
    MethodologyType.ISHIKAWA:     "Ishikawa (Рыбья кость)",
    MethodologyType.FTA:          "FTA (Дерево отказов)",
    MethodologyType.RCA_SYSTEMIC: "RCA Системный",
    MethodologyType.BOWTIE:       "Bowtie (Бабочка)",
}


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def generate_docx(result: RCAResult) -> bytes:
    """
    Генерирует DOCX-документ и возвращает его как bytes.
    Готов для StreamingResponse в FastAPI.
    """
    doc = Document()
    _setup_styles(doc)

    _add_header(doc, result)
    _add_summary(doc, result)

    if result.methodology == MethodologyType.BOWTIE:
        _add_bowtie_section(doc, result)
    else:
        _add_causal_sections(doc, result)

    _add_recommendations(doc, result)
    _add_meta(doc, result)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

def _setup_styles(doc: Document) -> None:
    """Настраивает поля страницы."""
    section = doc.sections[0]
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)


# ---------------------------------------------------------------------------
# Sections
# ---------------------------------------------------------------------------

def _add_header(doc: Document, result: RCAResult) -> None:
    methodology_label = METHODOLOGY_LABELS.get(result.methodology, result.methodology.value)
    created = result.created_at.strftime("%d.%m.%Y %H:%M UTC")

    # Заголовок
    h = doc.add_heading(level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.add_run(f"Отчёт RCA — {methodology_label}")
    run.font.color.rgb = _C_TITLE
    run.font.size = Pt(18)

    # Мета-строка
    p = doc.add_paragraph()
    _colored_run(p, f"ID: ", _C_MUTED, bold=False, size=9)
    _colored_run(p, result.result_id, _C_MUTED, bold=False, size=9)
    _colored_run(p, f"  |  Дата: ", _C_MUTED, bold=False, size=9)
    _colored_run(p, created, _C_MUTED, bold=False, size=9)
    _colored_run(p, f"  |  Модель: ", _C_MUTED, bold=False, size=9)
    _colored_run(p, result.model_used, _C_MUTED, bold=False, size=9)
    _colored_run(p, f"  |  Токены: ", _C_MUTED, bold=False, size=9)
    _colored_run(p, str(result.tokens_used), _C_MUTED, bold=False, size=9)
    _colored_run(p, f"  |  Уверенность: ", _C_MUTED, bold=False, size=9)
    _colored_run(p, f"{result.confidence_avg * 100:.0f}%", _C_MUTED, bold=False, size=9)

    doc.add_paragraph()  # отступ


def _add_summary(doc: Document, result: RCAResult) -> None:
    h = doc.add_heading(level=2)
    h.add_run("Резюме").font.color.rgb = _C_PRIMARY

    p = doc.add_paragraph(result.summary)
    p.runs[0].font.size = Pt(11)
    doc.add_paragraph()


def _add_causal_sections(doc: Document, result: RCAResult) -> None:
    """
    Универсальные секции для ishikawa / five_why / fta / rca_systemic.
    """
    sections = [
        ("Корневые причины",         result.root_causes,         _C_RED),
        ("Способствующие факторы",   result.contributing_causes, _C_ORANGE),
        ("Непосредственные причины", result.immediate_causes,    _C_BLUE),
    ]
    for title, nodes, color in sections:
        if not nodes:
            continue
        h = doc.add_heading(level=2)
        h.add_run(title).font.color.rgb = color

        for node in nodes:
            p = doc.add_paragraph(style="List Bullet")
            _colored_run(p, node.text, _C_TITLE, size=11)
            meta = f"  [{node.category}  {node.confidence * 100:.0f}%]"
            _colored_run(p, meta, _C_MUTED, bold=False, size=9)

        doc.add_paragraph()


def _add_bowtie_section(doc: Document, result: RCAResult) -> None:
    """
    Bowtie-специфичные секции, разделённые по category.
    category содержит BOWTIE:THREAT, BOWTIE:PREVENTION, BOWTIE:TOP_EVENT,
    BOWTIE:MITIGATION, BOWTIE:CONSEQUENCE, BOWTIE:HAZARD.
    """
    all_nodes = result.causal_tree

    def _by_cat(prefix: str):
        return [n for n in all_nodes if n.category.upper().startswith(f"BOWTIE:{prefix}")]

    hazards      = _by_cat("HAZARD")
    top_events   = _by_cat("TOP_EVENT")
    threats      = _by_cat("THREAT")
    prev_bars    = _by_cat("PREVENTION")
    consequences = _by_cat("CONSEQUENCE")
    miti_bars    = _by_cat("MITIGATION")

    # Fallback: если category не проставлен — используем стандартное отображение
    if not any([hazards, top_events, threats, prev_bars, consequences, miti_bars]):
        _add_causal_sections(doc, result)
        return

    bowtie_sections = [
        ("Опасный фактор (Hazard)",            hazards,      _C_MUTED),
        ("Топ-событие",                        top_events,   _C_RED),
        ("Угрозы",                             threats,      _C_RED),
        ("Барьеры предотвращения",             prev_bars,    _C_ORANGE),
        ("Последствия",                        consequences, _C_BLUE),
        ("Барьеры смягчения последствий",      miti_bars,    _C_GREEN),
    ]

    for title, nodes, color in bowtie_sections:
        if not nodes:
            continue
        h = doc.add_heading(level=2)
        h.add_run(title).font.color.rgb = color

        for node in nodes:
            p = doc.add_paragraph(style="List Bullet")
            _colored_run(p, node.text, _C_TITLE, size=11)
            # Деградированный барьер: confidence <= 0.3
            degraded = node.confidence <= 0.3 and "BARRIER" in node.category.upper()
            suffix = "  ⚠ деградирован" if degraded else ""
            meta = f"  [{node.confidence * 100:.0f}%{suffix}]"
            _colored_run(p, meta, _C_MUTED, bold=False, size=9)

        doc.add_paragraph()


def _add_recommendations(doc: Document, result: RCAResult) -> None:
    if not result.recommendations:
        return

    h = doc.add_heading(level=2)
    h.add_run(f"Рекомендации ({len(result.recommendations)})").font.color.rgb = _C_GREEN

    # Таблица: №, Приоритет, Категория, Ответственный, Текст
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for i, label in enumerate(["№", "Приоритет", "Категория", "Ответственный", "Рекомендация"]):
        hdr[i].text = label
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
        hdr[i].paragraphs[0].runs[0].font.color.rgb = _C_TITLE

    for idx, rec in enumerate(result.recommendations, start=1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = rec.priority.upper()
        row[2].text = rec.category
        row[3].text = rec.responsible or "—"
        row[4].text = rec.text

        # Цвет приоритета
        color = _PRIORITY_COLORS.get(rec.priority, _C_MUTED)
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
        # Красим ячейку приоритета
        for run in row[1].paragraphs[0].runs:
            run.font.color.rgb = color
            run.font.bold = True

    doc.add_paragraph()


def _add_meta(doc: Document, result: RCAResult) -> None:
    doc.add_heading(level=2).add_run("Техническая информация").font.color.rgb = _C_MUTED

    rows = [
        ("result_id",      result.result_id),
        ("incident_id",    result.incident_id),
        ("methodology",    result.methodology.value),
        ("model_used",     result.model_used),
        ("tokens_used",    str(result.tokens_used)),
        ("confidence_avg", f"{result.confidence_avg:.3f}"),
        ("created_at",     result.created_at.strftime("%Y-%m-%d %H:%M:%S UTC")),
        ("exported_at",    datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")),
    ]

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].text = k
        cells[1].text = v
        for para in cells[0].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = _C_MUTED
        for para in cells[1].paragraphs:
            for run in para.runs:
                run.font.size = Pt(8)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _colored_run(
    paragraph,
    text: str,
    color: RGBColor,
    bold: bool = True,
    size: int = 11,
) -> None:
    run = paragraph.add_run(text)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.size = Pt(size)
