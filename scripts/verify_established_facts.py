"""
Ручная проверка извлечения раздела «Установленные факты» из длинного DOCX.

Скрипт:
  1. Генерирует (или принимает) длинный .docx (~165 000 символов), где раздел
     «Установленные факты» расположен в середине документа — в той самой
     «мёртвой зоне», которую старая стратегия head+tail теряла.
  2. Прогоняет извлечение текста (extract_text_from_docx) и обрезку
     (_trim_text) — БЕЗ обращения к LLM.
  3. Проверяет, что раздел «Установленные факты» попал в срез.

Запуск:
    python scripts/verify_established_facts.py            # сгенерировать и проверить
    python scripts/verify_established_facts.py path.docx  # проверить свой документ
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

from src.services.docx_extractor import extract_text_from_docx
from src.services.docx_fields_service import _trim_text

FACTS_MARKER = "Установленные факты"
FACTS_BODY = (
    "Установленные факты: комиссия по расследованию несчастного случая "
    "установила, что причиной послужило нарушение технологического регламента "
    "и отсутствие надлежащего контроля со стороны ответственных лиц. "
)


def build_sample_docx(path: Path) -> None:
    """Создаёт длинный документ с разделом «Установленные факты» в середине."""
    doc = Document()
    doc.add_heading("Акт о расследовании несчастного случая", level=1)
    doc.add_paragraph("Раздел 1. Общие сведения об организации и пострадавших.")

    # Наполнитель ДО раздела (имитация большого тела документа)
    filler = "Описание обстоятельств, нормативных требований и хронологии. " * 20
    for _ in range(40):
        doc.add_paragraph(filler)

    # Ключевой раздел — в СЕРЕДИНЕ документа
    doc.add_heading("Установленные факты", level=2)
    for _ in range(8):
        doc.add_paragraph(FACTS_BODY)

    # Наполнитель ПОСЛЕ раздела
    for _ in range(40):
        doc.add_paragraph(filler)

    doc.save(path)


def main() -> int:
    if len(sys.argv) > 1:
        path = Path(sys.argv[1])
        if not path.exists():
            print(f"Файл не найден: {path}")
            return 2
    else:
        path = Path("/tmp/sample_incident.docx")
        build_sample_docx(path)
        print(f"Сгенерирован тестовый документ: {path}")

    file_bytes = path.read_bytes()
    text = extract_text_from_docx(file_bytes)
    print(f"Извлечено символов из DOCX: {len(text)}")

    facts_pos = text.lower().find(FACTS_MARKER.lower())
    print(f"Позиция раздела «{FACTS_MARKER}» в полном тексте: {facts_pos}")

    trimmed, was_trimmed = _trim_text(text)
    print(f"Обрезка применена: {was_trimmed}; длина среза: {len(trimmed)}")

    in_trimmed = FACTS_MARKER.lower() in trimmed.lower()
    print(f"Раздел «{FACTS_MARKER}» присутствует в срезе: {in_trimmed}")

    if in_trimmed:
        print("\n✅ УСПЕХ: раздел «Установленные факты» попадёт в запрос к LLM.")
        return 0

    print("\n❌ ПРОВАЛ: раздел не попал в срез — established_facts останется пустым.")
    return 1


if __name__ == "__main__":
    raise SystemExit(main())
